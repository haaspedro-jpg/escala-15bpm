import streamlit as st
import pandas as pd
import random
import json
import os
import time
import requests
import re
import unicodedata
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT

# --- CONFIGURAÇÕES DO BANCO DE DADOS NA NUVEM (JSONBIN.IO) ---
BIN_ID = "6a0dc2716877513b27a34d82"
API_KEY = "$2a$10$mtrG79e1mc862Rn3ZbwgZu7hJhbdHyTb.bZqvA7cmvNKyxM9gOG9G"
URL_DB = f"https://api.jsonbin.io/v3/b/{BIN_ID}"
HEADERS_DB = {
    'Content-Type': 'application/json',
    'X-Master-Key': API_KEY
}

# --- DICIONÁRIO OFICIAL DE NOMES DE GUERRA ---
NOMES_DE_GUERRA = {
    "Alan Felipe Buck dos Santos": "FELIPE",
    "Ana Paula Machado": "MACHADO",
    "Arthur Amador Silva": "ARTHUR",
    "Carina dos Santos": "CARINA",
    "Carlos Eduardo Boava": "CARLOS BOAVA",
    "Carlos Eduardo Soares da Costa Filho": "DA COSTA",
    "Cláudio Márcio Silva Júnior": "CLÁUDIO",
    "Daniel Barbosa Ramos": "RAMOS",
    "Danilo Eduardo de Oliveira": "OLIVEIRA",
    "Dener Vinícius Gomes Rodrigo": "DENER",
    "Edilma Aparecida Sereia da Silva": "SEREIA",
    "Eduardo Cauan José": "JOSÉ",
    "Eduardo Júnior Marques": "JÚNIOR MARQUES",
    "Felipe Israel de Almeida Guilherme": "ISRAEL",
    "Filipe Aparecido dos Santos": "SANTOS",
    "Gabriel Magdiel Reis": "MAGDIEL",
    "Gabriela Regina Schneider": "SCHNEIDER",
    "Gustavo Henrique Nogueira": "NOGUEIRA",
    "Gustavo Tivirolli Favaro": "TIVIROLLI",
    "Hemerson Oliveira Pacheco Júnior": "HEMERSON",
    "Hilda Heloisa Andrade Cunha": "ANDRADE",
    "Igor Henrique Bibanco": "BIBANCO",
    "Israel Mendes Martins": "MENDES",
    "João Victor Bertola da Silva": "BERTOLA",
    "João Victor Rapharred de Oliveira Abrahao": "RAPHARRED",
    "Jonathan Henrique das Neves": "HENRIQUE",
    "Josué Almeida de Souza": "JOSUÉ",
    "Kauane Stefany Lopes": "LOPES",
    "Leonardo Brenes Burque": "BURQUE",
    "Leonardo Cardoso Cosmos": "COSMOS",
    "Leonardo Santiago Rodrigues": "SANTIAGO",
    "Letícia Lopes Martelossi": "LETÍCIA",
    "Lirian Borges Barbosa": "LIRIAN",
    "Lucas Eduardo Dornellas de Oliveira": "EDUARDO",
    "Lucas Gabriel Souza de Campos": "CAMPOS",
    "Lucas Matheus dos Santos Lopes": "LUCAS LOPES",
    "Lucas Silva Piccioni": "PICCIONI",
    "Lucas Vinícius da Silva Panta": "PANTA",
    "Mateus Fujita Araújo": "FUJITA",
    "Matheus Silva Aleluia": "ALELUIA",
    "Nathalia Vargas Lopes": "VARGAS",
    "Priscila de Andrade Correia": "CORREIA",
    "Rafael Ribeiro Garcia": "RAFAEL GARCIA",
    "Ricardo Vinícius Zamparoni": "ZAMPARONI",
    "Tailer Costa Ribeiro dos Santos": "TAILER",
    "Thiago Zanin": "THIAGO ZANIN",
    "Thialis Venâncio dos Santos": "VENÂNCIO",
    "Tiago da Silva Ferreira": "TIAGO FERREIRA",
    "Victor Gonçalves Brum Silva": "BRUM",
    "Vitor Hugo Cardoso Almondes": "ALMONDES",
    "Vitor Humberto Ortega Nascimento": "ORTEGA",
    "Vitor Kauan Amorim dos Reis": "VITOR REIS",
    "Vitoria Cristina Cortelini": "CORTELINI",
    "Wallison Rodrigo Passerini": "PASSERINI",
    "Wesley Rodrigues Padovan": "PADOVAN"
}

MESES_PT = {1: 'JANEIRO', 2: 'FEVEREIRO', 3: 'MARÇO', 4: 'ABRIL', 5: 'MAIO', 6: 'JUNHO', 
            7: 'JULHO', 8: 'AGOSTO', 9: 'SETEMBRO', 10: 'OUTUBRO', 11: 'NOVEMBRO', 12: 'DEZEMBRO'}
DIAS_SEMANA_PT = {0: 'segunda', 1: 'terça', 2: 'quarta', 3: 'quinta', 4: 'sexta', 5: 'sábado', 6: 'domingo'}

@st.cache_data(ttl=5)
def carregar_banco_nuvem():
    try:
        req = requests.get(URL_DB, headers=HEADERS_DB)
        if req.status_code == 200:
            data = req.json()['record']
            if 'ativo' not in data:
                data = {'ativo': {}, 'backups': {}, 'escalas_arquivadas': {}}
            if 'escalas_arquivadas' not in data:
                data['escalas_arquivadas'] = {}
            return data
        else:
            st.error("Falha ao conectar com o banco de dados da nuvem.")
            return None
    except Exception as e:
        st.error(f"Erro de rede: {e}")
        return None

def salvar_banco_nuvem(dados_completos):
    try:
        req = requests.put(URL_DB, json=dados_completos, headers=HEADERS_DB)
        if req.status_code == 200:
            st.cache_data.clear()
            return True
        return False
    except:
        return False

def ler_planilha_segura(caminho_ou_arquivo):
    try:
        if str(caminho_ou_arquivo).endswith('.xlsx'):
            df = pd.read_excel(caminho_ou_arquivo)
        else:
            try:
                df = pd.read_csv(caminho_ou_arquivo, sep=None, engine='python', encoding='utf-8')
            except UnicodeDecodeError:
                df = pd.read_csv(caminho_ou_arquivo, sep=None, engine='python', encoding='latin-1')
        df.columns = df.columns.str.strip().str.upper()
        return df
    except Exception:
        return None

def formatar_texto_pm(run, font_name='Arial', size_pt=10, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold

def estilizar_celula(cell, texto, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = str(texto)
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0
    for run in p.runs:
        formatar_texto_pm(run, 'Arial', 10, bold=bold)

def obter_nome_guerra(nome_completo):
    def clean(s):
        return ''.join(c for c in unicodedata.normalize('NFKD', s) if not unicodedata.combining(c)).upper().strip()
    
    nome_clean = clean(nome_completo)
    for original, guerra in NOMES_DE_GUERRA.items():
        if clean(original) == nome_clean:
            return guerra
    return None

def estilizar_celula_nome(cell, nome_completo, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.0

    nome_guerra = obter_nome_guerra(nome_completo)
    
    if not nome_guerra:
        run = p.add_run(nome_completo)
        formatar_texto_pm(run, 'Arial', 10, bold=False)
        return

    def clean(s):
        return ''.join(c for c in unicodedata.normalize('NFKD', s) if not unicodedata.combining(c)).upper()

    guerra_words = set(clean(nome_guerra).split())
    parts = re.split(r'(\s+)', nome_completo)

    for part in parts:
        if not part.strip():
            run = p.add_run(part)
            formatar_texto_pm(run, 'Arial', 10, bold=False)
        else:
            if clean(part) in guerra_words:
                run = p.add_run(part)
                formatar_texto_pm(run, 'Arial', 10, bold=True)
            else:
                run = p.add_run(part)
                formatar_texto_pm(run, 'Arial', 10, bold=False)

def gerar_documento_bytes(escala_gerada):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.15
    
    texto_cabecalho = (
        "ESTADO DO PARANÁ\n"
        "POLÍCIA MILITAR\n"
        "2º COMANDO REGIONAL DE POLÍCIA MILITAR\n"
        "NÚCLEO DE ENSINO DO 15º BPM\n"
        "CFP 15º BPM 2026/2027\n\n"
        "ESCALA DE SERVIÇO DE GUARDA - SEDE DO 15ºBPM\n"
        "Horário 07h00min. às 07h00min.\n"
    )
    run_cabecalho = p.add_run(texto_cabecalho)
    formatar_texto_pm(run_cabecalho, 'Arial', 10, bold=True)
    
    for dia_str, plantonistas in escala_gerada.items():
        p_dia = doc.add_paragraph()
        p_dia.paragraph_format.space_before = Pt(12)
        p_dia.paragraph_format.space_after = Pt(4)
        run_dia = p_dia.add_run(dia_str)
        formatar_texto_pm(run_dia, 'Arial', 10, bold=True)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        hdr_cells = table.rows[0].cells
        headers = ['QUARTO\nDE HORA', 'GRAD.', 'NOME', 'CPF']
        for idx, text in enumerate(headers):
            estilizar_celula(hdr_cells[idx], text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            hdr_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        quartos = ['1º', '2º', '3º']
        for idx in range(3):
            militar1 = plantonistas[idx * 2]
            militar2 = plantonistas[idx * 2 + 1]
            
            row1 = table.add_row().cells
            estilizar_celula(row1[0], quartos[idx], bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            estilizar_celula(row1[1], 'Sd. 3ª Classe PM', bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
            estilizar_celula_nome(row1[2], militar1['nome'], align=WD_ALIGN_PARAGRAPH.LEFT)
            estilizar_celula(row1[3], militar1['cpf'], bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
            
            row2 = table.add_row().cells
            estilizar_celula(row2[0], '', bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
            estilizar_celula(row2[1], 'Sd. 3ª Classe PM', bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
            estilizar_celula_nome(row2[2], militar2['nome'], align=WD_ALIGN_PARAGRAPH.LEFT)
            estilizar_celula(row2[3], militar2['cpf'], bold=False, align=WD_ALIGN_PARAGRAPH.CENTER)
            
            cell_qto_1 = row1[0]
            cell_qto_2 = row2[0]
            cell_qto_1.merge(cell_qto_2)
            cell_qto_1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
        for row in table.rows:
            row.cells[0].width = Cm(2.5)
            row.cells[1].width = Cm(3.5)
            row.cells[2].width = Cm(8.0)
            row.cells[3].width = Cm(3.0)

    p_ass = doc.add_paragraph()
    p_ass.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ass.paragraph_format.space_before = Pt(24)
    run_ass = p_ass.add_run("\n\nAssinado Eletronicamente\n1º Ten. QOEM PM Jeferson Luis Fracaro,\nCoordenador do CFP 15º BPM.")
    formatar_texto_pm(run_ass, 'Arial', 10, bold=True)
    
    import io
    target_stream = io.BytesIO()
    doc.save(target_stream)
    return target_stream.getvalue()

feriados_nacionais_2026 = [
    datetime(2026, 1, 1).date(), datetime(2026, 2, 17).date(), datetime(2026, 4, 3).date(),
    datetime(2026, 4, 21).date(), datetime(2026, 5, 1).date(), datetime(2026, 6, 4).date(),
    datetime(2026, 9, 7).date(), datetime(2026, 10, 12).date(), datetime(2026, 11, 2).date(),
    datetime(2026, 11, 15).date(), datetime(2026, 12, 25).date()
]

# --- INTERFACE STREAMLIT ---
st.set_page_config(page_title="Gerador de Escala CFP", layout="wide")
st.title("Gerador Automático - Escala de Guarda CFP")

if 'escala_recem_gerada' not in st.session_state:
    st.session_state['escala_recem_gerada'] = None
if 'nome_arquivo_recem_gerado' not in st.session_state:
    st.session_state['nome_arquivo_recem_gerado'] = None

# Agora prioriza ler o arquivo .xlsx
caminho_planilha = 'alunos.xlsx' if os.path.exists('alunos.xlsx') else 'alunos.csv'

if not os.path.exists(caminho_planilha):
    st.error("O arquivo 'alunos.xlsx' ou 'alunos.csv' não foi encontrado. Suba a planilha no repositório.")
    st.stop()

alunos_df = ler_planilha_segura(caminho_planilha)

if alunos_df is None:
    st.error("Erro ao ler a planilha. Certifique-se de que ela está salva corretamente.")
    st.stop()

if 'NOME' not in alunos_df.columns or 'CPF' not in alunos_df.columns:
    st.error("Erro: A planilha precisa ter as colunas exatas 'NOME' e 'CPF'.")
    st.stop()
    
dados_nuvem = carregar_banco_nuvem()
if dados_nuvem is None:
    st.stop()
    
historico_ativo = dados_nuvem.get('ativo', {})
backups_nuvem = dados_nuvem.get('backups', {})
escalas_arquivadas = dados_nuvem.get('escalas_arquivadas', {})

# Sincronização automática inicial da planilha (Conserta os nomes defeituosos)
precisa_salvar = False
for index, row in alunos_df.iterrows():
    cpf = str(row['CPF']).strip()
    nome = str(row['NOME']).strip()
    if cpf not in historico_ativo:
        historico_ativo[cpf] = {'nome': nome, 'normal': 0, 'vermelho': 0, 'ultimo_servico': None, 'datas_servico': []}
        precisa_salvar = True
    else:
        # Se o nome no banco tiver símbolos quebrados e a planilha nova estiver certa, ele corrige sozinho
        if historico_ativo[cpf]['nome'] != nome:
            historico_ativo[cpf]['nome'] = nome
            precisa_salvar = True
            
if precisa_salvar:
    dados_nuvem['ativo'] = historico_ativo
    salvar_banco_nuvem(dados_nuvem)

# CONTROLES E ABAS DE VISUALIZAÇÃO
with st.expander("📊 PAINEL DE AUDITORIA E CONSULTA DE ESCALAS", expanded=False):
    tab_geral, tab_individual, tab_config = st.tabs(["📋 Visão Geral (Saldos)", "🔍 Consulta Individualizada por Aluno", "⚙️ Configuração Inicial"])
    
    with tab_geral:
        dados_tabela = []
        for cpf, d in historico_ativo.items():
            datas = d.get('datas_servico', [])
            ultimo_str = f"{datas[-1]['data']} ({datas[-1]['tipo']})" if datas else "Nenhum dia registrado"
            dados_tabela.append({
                'Nome do Aluno': d['nome'],
                'Normais': d['normal'],
                'Vermelhas': d['vermelho'],
                'Total': d['normal'] + d['vermelho'],
                'Último Serviço': ultimo_str
            })
        df_historico = pd.DataFrame(dados_tabela)
        df_historico = df_historico.sort_values(by=['Vermelhas', 'Normais', 'Nome do Aluno'])
        st.dataframe(df_historico, use_container_width=True)
        
    with tab_individual:
        nomes_alunos = sorted([d['nome'] for d in historico_ativo.values()])
        aluno_sel = st.selectbox("Selecione o Aluno-Soldado para detalhamento:", nomes_alunos)
        cpf_sel = [cpf for cpf, d in historico_ativo.items() if d['nome'] == aluno_sel][0]
        dados_aluno = historico_ativo[cpf_sel]
        
        st.markdown(f"### Ficha de Serviço: {dados_aluno['nome']}")
        col_b1, col_b2, col_b3 = st.columns(3)
        col_b1.metric("Escalas em Dias Normais", dados_aluno['normal'])
        col_b2.metric("Escalas em Dias Vermelhos", dados_aluno['vermelho'])
        col_b3.metric("Total Geral de Escalas", dados_aluno['normal'] + dados_aluno['vermelho'])
        
        if dados_aluno.get('datas_servico'):
            df_individual = pd.DataFrame(dados_aluno['datas_servico'])
            df_individual.columns = ['📅 Dia da Escala', '🏷️ Tipo de Dia']
            df_individual = df_individual.sort_values(by='📅 Dia da Escala', ascending=False)
            st.table(df_individual)
        else:
            st.info("Este militar ainda não possui nenhuma escala registrada.")

    with tab_config:
        st.markdown("### Configurar Histórico Inicial (Ponto Zero)")
        st.write("Clique no botão abaixo apenas se o sistema estiver zerado para injetar o histórico das duas escalas antigas (11 a 24 de Maio).")
        if st.button("🔄 Inicializar Ponto Zero (Escalas de 11 a 24 de Maio)"):
            servicos_normais = ["Gabriela Regina Schneider", "Filipe Aparecido dos Santos", "Lucas Gabriel Souza de Campos", "João Victor Rapharred de Oliveira Abrahao", "Kauane Stefany Lopes", "Vitor Humberto Ortega Nascimento", "Lucas Eduardo Dornellas de Oliveira", "Tailer Costa Ribeiro dos Santos", "Daniel Barbosa Ramos", "Lorenna Elen de Souza Chiconato", "Carlos Eduardo Boava", "Ricardo Vinícius Zamparoni", "Lirian Borges Barbosa", "Wesley Rodrigues Padovan", "Victor Gonçalves Brum Silva", "Jonathan Henrique das Neves", "Thialis Venâncio dos Santos", "Rafael Ribeiro Garcia", "Gustavo Tivirolli Favaro", "Vitor Hugo Cardoso Almondes", "Danilo Eduardo de Oliveira", "Wallison Rodrigo Passerini", "Matheus Silva Aleluia", "Gustavo Henrique Nogueira", "Igor Henrique Bibanco", "Lucas Matheus dos Santos Lopes", "Carlos Eduardo Soares da Costa Filho", "Thiago Zanin", "Dener Vinícius Gomes Rodrigo", "Cláudio Márcio Silva Júnior", "Nathalia Vargas Lopes", "Felipe Israel de Almeida Guilherme", "Josué Almeida de Souza", "Leonardo Santiago Rodrigues", "Hemerson Oliveira Pacheco Junior", "Ana Paula Machado", "Tiago da Silva Ferreira", "Leonardo Brenes Burque", "Gabriel Magdiel Reis", "Letícia Lopes Martelossi", "Vitor Hugo Cardoso Almondes", "João Victor Bertola da Silva", "Lucas Vinicius da Silva Panta", "Eduardo Junior Marques", "Gabriela Regina Schneider", "Lucas Eduardo Dornellas de Oliveira", "Filipe Aparecido dos Santos", "Victor Gonçalves Brum Silva", "Jonathan Henrique das Neves", "Thialis Venâncio dos Santos", "Rafael Ribeiro Garcia", "Carina dos Santos", "Tailer Costa Ribeiro dos Santos", "Kauane Stefany Lopes", "Daniel Barbosa Ramos", "Vitor Humberto Ortega Nascimento", "João Victor Rapharred de Oliveira Abrahao", "Lucas Gabriel Souza de Campos", "Wesley Rodrigues Padovan", "Lirian Borges Barbosa"]
            servicos_vermelhos = ["Hilda Heloisa Andrade Cunha", "Vitoria Cristina Cortelini", "Priscila de Andrade Correia", "Mateus Fujita Araújo", "Lucas Silva Piccioni", "Leonardo Cardoso Cosmos", "Alan Felipe Buck dos Santos", "Arthur Amador Silva", "Vitor Kauan Amorim dos Reis", "Edilma Aparecida Sereia da Silva", "Israel Mendes Martins", "Eduardo Cauan José", "Igor Henrique Bibanco", "Lucas Matheus dos Santos Lopes", "Carlos Eduardo Soares da Costa Filho", "Gustavo Tivirolli Favaro", "Carlos Eduardo Boava", "Ricardo Vinícius Zamparoni", "Vitor Hugo Cardoso Almondes", "Danilo Eduardo de Oliveira", "Wallison Rodrigo Passerini", "Matheus Silva Aleluia", "Gustavo Henrique Nogueira", "Thiago Zanin"]
            
            for cpf, dados in historico_ativo.items():
                n_aluno = dados['nome'].lower().strip()
                qtd_normal = sum(1 for n in servicos_normais if n.lower().strip() in n_aluno)
                qtd_vermelho = sum(1 for v in servicos_vermelhos if v.lower().strip() in n_aluno)
                
                historico_ativo[cpf]['normal'] = qtd_normal
                historico_ativo[cpf]['vermelho'] = qtd_vermelho
                historico_ativo[cpf]['datas_servico'] = []
                for _ in range(qtd_normal):
                    historico_ativo[cpf]['datas_servico'].append({"data": "Maio/2026 (Base)", "tipo": "Normal"})
                for _ in range(qtd_vermelho):
                    historico_ativo[cpf]['datas_servico'].append({"data": "Maio/2026 (Base)", "tipo": "Vermelha"})
            
            dados_nuvem['ativo'] = historico_ativo
            dados_nuvem['backups']["backup_PONTO_INICIAL_BASE"] = json.loads(json.dumps(historico_ativo))
            if salvar_banco_nuvem(dados_nuvem):
                st.success("Ponto Zero configurado com sucesso e salvo na nuvem!")
                time.sleep(1)
                st.rerun()

# ACERVO DE DOCUMENTOS SALVOS NA NUVEM
with st.expander("📂 ARQUIVOS DE ESCALAS JÁ GERADAS (Disponíveis para Download)", expanded=False):
    if escalas_arquivadas:
        for nome_arq, dados_escala in sorted(escalas_arquivadas.items(), reverse=True):
            doc_bytes = gerar_documento_bytes(dados_escala)
            st.download_button(
                label=f"📥 Baixar {nome_arq}",
                data=doc_bytes,
                file_name=nome_arq,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=nome_arq
            )
    else:
        st.info("Nenhuma escala arquivada no banco de dados ainda.")

st.divider()

st.subheader("Configuração do Período")
col1, col2 = st.columns(2)
with col1:
    data_inicio = st.date_input("Data de Início", format="DD/MM/YYYY")
with col2:
    data_fim = st.date_input("Data de Término", format="DD/MM/YYYY")
    
feriados = st.text_input("Feriados MUNICIPAIS (opcional - separados por vírgula DD/MM/AAAA):")
feriados_lista = []
if feriados:
    try:
        feriados_lista = [datetime.strptime(f.strip(), "%d/%m/%Y").date() for f in feriados.split(",") if f.strip()]
    except ValueError:
        st.warning("Formato de feriado inválido.")
        
todos_feriados = feriados_nacionais_2026 + feriados_lista

if st.button("Gerar Escala", type="primary"):
    if data_inicio > data_fim:
        st.error("A data de início deve ser anterior ou igual à data de término.")
    else:
        dias_ja_gerados = set(item['data'] for dados in historico_ativo.values() for item in dados.get('datas_servico', []))
        conflitos = []
        delta = data_fim - data_inicio
        for i in range(delta.days + 1):
            dia_atual = data_inicio + timedelta(days=i)
            if dia_atual.strftime('%d/%m/%Y') in dias_ja_gerados:
                conflitos.append(dia_atual.strftime('%d/%m/%Y'))
                
        if conflitos:
            st.error(f"⚠️ **BLOQUEADO:** Já existe escala para: **{', '.join(conflitos)}**.")
            st.stop()

        nome_backup = f"Backup_antes_de_{data_inicio.strftime('%d_%m')}_a_{data_fim.strftime('%d_%m')}_{datetime.now().strftime('%H%M%S')}"
        dados_nuvem['backups'][nome_backup] = json.loads(json.dumps(historico_ativo))

        escala_final = {}
        historico_temp = json.loads(json.dumps(historico_ativo))
        
        for i in range(delta.days + 1):
            dia_atual = data_inicio + timedelta(days=i)
            dia_semana = dia_atual.weekday()
            
            is_vermelho = dia_semana >= 5 or dia_atual in todos_feriados
            tipo_dia = 'vermelho' if is_vermelho else 'normal'
            
            texto_data = f"EFETIVO: DIA {dia_atual.day:02d} DE {MESES_PT[dia_atual.month]} DE {dia_atual.year} – Das 07h00 de {DIAS_SEMANA_PT[dia_semana]} às 07h00 de {DIAS_SEMANA_PT[(dia_atual + timedelta(days=1)).weekday()]} "
            data_formatada_simples = dia_atual.strftime('%d/%m/%Y')
            
            cpf_elegiveis = []
            for cpf, dados in historico_temp.items():
                elegivel = True
                if dados.get('ultimo_servico'):
                    ultimo_servico_data = datetime.strptime(dados['ultimo_servico'], '%Y-%m-%d').date()
                    if (dia_atual - ultimo_servico_data).days <= 1:
                        elegivel = False
                        
                if dia_semana == 0 and "haas" in dados['nome'].lower():
                    elegivel = False
                    
                if elegivel:
                    cpf_elegiveis.append(cpf)
            
            if len(cpf_elegiveis) < 6:
                st.error(f"Erro: Efetivo insuficiente para o dia {data_formatada_simples}.")
                st.stop()
                
            cpf_elegiveis.sort(key=lambda cpf: (
                historico_temp[cpf][tipo_dia], 
                historico_temp[cpf]['normal'] + historico_temp[cpf]['vermelho'], 
                random.random()
            ))
            
            selecionados = cpf_elegiveis[:6]
            plantonistas_do_dia = []
            
            for cpf in selecionados:
                historico_temp[cpf][tipo_dia] += 1
                historico_temp[cpf]['ultimo_servico'] = dia_atual.strftime('%Y-%m-%d')
                if 'datas_servico' not in historico_temp[cpf]:
                    historico_temp[cpf]['datas_servico'] = []
                
                historico_temp[cpf]['datas_servico'].append({
                    "data": data_formatada_simples,
                    "tipo": "Vermelha" if tipo_dia == "vermelho" else "Normal"
                })
                plantonistas_do_dia.append({'cpf': cpf, 'nome': historico_temp[cpf]['nome']})
                
            escala_final[texto_data] = plantonistas_do_dia

        nome_arquivo_novo = f"escala_gerada_{data_inicio.strftime('%d_%m')}_a_{data_fim.strftime('%d_%m')}.docx"
        
        dados_nuvem['escalas_arquivadas'][nome_arquivo_novo] = escala_final
        dados_nuvem['ativo'] = historico_temp
        
        if salvar_banco_nuvem(dados_nuvem):
            st.session_state['escala_recem_gerada'] = escala_final
            st.session_state['nome_arquivo_recem_gerado'] = nome_arquivo_novo
            st.rerun()

if st.session_state['escala_recem_gerada']:
    st.success("Escala gerada com sucesso!")
    bytes_word = gerar_documento_bytes(st.session_state['escala_recem_gerada'])
    st.download_button(
        label="📥 Baixar Documento Recém Gerado (.docx)",
        data=bytes_word,
        file_name=st.session_state['nome_arquivo_recem_gerado'],
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

st.divider()
st.subheader("⚠️ Linha do Tempo de Backups e Restauração")
if backups_nuvem:
    opcoes_backup = sorted(list(backups_nuvem.keys()), reverse=True)
    backup_selecionado = st.selectbox("Selecione um ponto de restauração para desfazer um sorteio:", options=opcoes_backup)
    
    if st.button("Restaurar para o Ponto Selecionado", type="secondary"):
        dados_nuvem['ativo'] = dados_nuvem['backups'][backup_selecionado]
        chaves_escalas = sorted(list(dados_nuvem['escalas_arquivadas'].keys()))
        if chaves_escalas:
            del dados_nuvem['escalas_arquivadas'][chaves_escalas[-1]]
        
        del dados_nuvem['backups'][backup_selecionado]
        if salvar_banco_nuvem(dados_nuvem):
            st.session_state['escala_recem_gerada'] = None
            st.success("Sistema restaurado com sucesso na nuvem!")
            time.sleep(1)
            st.rerun()
else:
    st.info("Nenhum ponto de restauração criado na nuvem ainda.")
